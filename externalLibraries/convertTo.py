from docx import Document
import time
from docxtpl import DocxTemplate, InlineImage
import os
import win32com.client
from docxtpl import DocxTemplate, InlineImage




def createWord(data,documentName,failTubes,goodTubes):

    #tomando el nombre del documento
    
    
    # Create a new document from template
    templateDoc=DocxTemplate("documentation/template2.docx")
    #variable que guarda la fecha
    fecha=time.strftime("%d/%m/%Y")
    #variable que guarda la hora
    hora=time.strftime("%H:%M:%S")
    #variable que guarda el nombre del test
    test="Test : "+documentName

    dataFormateada=""
    for i in range(len(data)):
        for j in range(len(data[i])):
            dataFormateada+='\t'+data[i][j]+"\t"
            
        dataFormateada+="\n-----------------------------------------------------------------------------------------------------------------\n"

    data=dataFormateada
    
    # Create a context for rendering
    context = {
        'fecha': fecha,
        'hora': hora,
        'test': test,
        'data': data,
        'failTubes': failTubes,
        'goodTubes': goodTubes,
        'totalTubes': failTubes+goodTubes
    }
    #render the document
    templateDoc.render(context)
    #save the document
    templateDoc.save("reportPDF/"+documentName+".docx")
    
    return documentName


def convertToPdf(documentName):
    wdFormatPDF = 17

    inputFile = os.path.abspath("reportPDF/"+documentName+".docx")
    outputFile = os.path.abspath("reportPDF/"+documentName+".pdf")
    
    word = win32com.client.Dispatch('Word.Application')
    try:
        doc = word.Documents.Open(inputFile)
        doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
        doc.Close()
        print("Conversión completada con éxito.")
    except Exception as e:
        print("Error al convertir el documento:", e)
    finally:
        word.Quit()


import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
import os 

def recipesExcel(report,path):
    # creo un libro de trabajo
    wb = Workbook()
    # creo una hoja de trabajo
    ws = wb.active
    # para cada lista voy a tener una hoja

    # creo una tabla con los datos de la lista
    table = Table(displayName="Recipes", ref="A1:k"+str(len(report)))
    # creo un estilo para la tabla
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                        showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    # aplico el estilo a la tabla
    table.tableStyleInfo = style
    # agrego la tabla a la hoja
    ws.add_table(table)
    # recorro la lista y la agrego a la hoja
    for i in range(len(report)):
        ws.cell(row=i+1, column=1).value = report[i][0]
        ws.cell(row=i+1, column=2).value = report[i][1]
        ws.cell(row=i+1, column=3).value = report[i][2]
        ws.cell(row=i+1, column=4).value = report[i][3]
        ws.cell(row=i+1, column=5).value = report[i][4]
        ws.cell(row=i+1, column=6).value = report[i][5]
        ws.cell(row=i+1, column=7).value = report[i][6]
        ws.cell(row=i+1, column=8).value = report[i][7]
        ws.cell(row=i+1, column=9).value = report[i][8]
        ws.cell(row=i+1, column=10).value = report[i][9]
        ws.cell(row=i+1, column=11).value = report[i][10]


    # creo un estilo para las celdas
    font = Font(name='Calibri', size=11, bold=False, italic=False, vertAlign=None, underline='none', strike=False, color='FF000000')

    # recorro la hoja y le aplico el estilo a las celdas
    for row in ws.iter_rows():
        for cell in row:
            cell.font = font

    # recorro la lista del reporte por fila y creo otro reporte contando cuantos yes or no por fila
    contadorTotalList=[["Description","Yes","No"]]
    for i in range(1,len(report)):
        #agrego la descripcion de cada fila
        contadorTotalList.append([report[i][1]])
        #agrego el contador de yes
        contadorTotalList[i].append(report[i].count("Yes"))
        #agrego el contador de no
        contadorTotalList[i].append(report[i].count("No"))
    
    # creo una hoja
    ws2 = wb.create_sheet("Total_Yes_No")
    # creo una tabla con los datos de la lista
    table = Table(displayName="Table1", ref="A1:C"+str(len(contadorTotalList)))
    # creo un estilo para la tabla
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                        showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    # aplico el estilo a la tabla
    table.tableStyleInfo = style
    # agrego la tabla a la hoja
    ws2.add_table(table)
    # recorro la lista y la agrego a la hoja
    for i in range(len(contadorTotalList)):
        ws2.cell(row=i+1, column=1).value = contadorTotalList[i][0]
        ws2.cell(row=i+1, column=2).value = contadorTotalList[i][1]
        ws2.cell(row=i+1, column=3).value = contadorTotalList[i][2]

    # creo un estilo para las celdas
    font = Font(name='Calibri', size=11, bold=False, italic=False, vertAlign=None, underline='none', strike=False, color='FF000000')

    # recorro la hoja y le aplico el estilo a las celdas
    for row in ws2.iter_rows():
        for cell in row:
            cell.font = font    
    
    
    wb.save(path)
    




# creo una funcion que reciba una lista de listas y cree un excel con esos datos
def createExcel(contadorTotalList,equiposFallandolist,tubosdesconectados,path):

    # creo un libro de trabajo
    wb = Workbook()
    # creo una hoja de trabajo
    ws = wb.active
    # para cada lista voy a tener una hoja

    # creo una tabla con los datos de la lista
    table = Table(displayName="Table1", ref="A1:C"+str(len(contadorTotalList)))
    # creo un estilo para la tabla
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                        showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    # aplico el estilo a la tabla
    table.tableStyleInfo = style
    # agrego la tabla a la hoja
    ws.add_table(table)
    # recorro la lista y la agrego a la hoja
    for i in range(len(contadorTotalList)):
        ws.cell(row=i+1, column=1).value = contadorTotalList[i][0]
        ws.cell(row=i+1, column=2).value = contadorTotalList[i][1]
        ws.cell(row=i+1, column=3).value = contadorTotalList[i][2]

    # creo un estilo para las celdas
    font = Font(name='Calibri', size=11, bold=False, italic=False, vertAlign=None, underline='none', strike=False, color='FF000000')

    # recorro la hoja y le aplico el estilo a las celdas
    for row in ws.iter_rows():
        for cell in row:
            cell.font = font
    
    # creo otra hoja
    ws2 = wb.create_sheet("Equipos_Desconectados")
    # creo una tabla con los datos de la lista
    table = Table(displayName="Table2", ref="A1:C"+str(len(equiposFallandolist)))
    # creo un estilo para la tabla y la amplio
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                        showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    
   

    # aplico el estilo a la tabla
    table.tableStyleInfo = style
    # agrego la tabla a la hoja
    ws2.add_table(table)
    # recorro la lista y la agrego a la hoja
    for i in range(len(equiposFallandolist)):
        ws2.cell(row=i+1, column=1).value = equiposFallandolist[i][0]
        ws2.cell(row=i+1, column=2).value = equiposFallandolist[i][1]
        ws2.cell(row=i+1, column=3).value = equiposFallandolist[i][2]
    
    # creo un estilo para las celdas
    font = Font(name='Calibri', size=11, bold=False, italic=False, vertAlign=None, underline='none', strike=False, color='FF000000')

    # recorro la hoja y le aplico el estilo a las celdas
    for row in ws2.iter_rows():
        for cell in row:
            cell.font = font


    # creo otra hoja

    ws3 = wb.create_sheet("Total_TipoEquipos")

    # creo una tabla con los datos de la lista

    table = Table(displayName="Table3", ref="A1:C"+str(len(tubosdesconectados)))
    
    # creo un estilo para la tabla

    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                        showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    
    # aplico el estilo a la tabla

    table.tableStyleInfo = style

    # agrego la tabla a la hoja

    ws3.add_table(table)

    # recorro la lista y la agrego a la hoja

    for i in range(len(tubosdesconectados)):
        ws3.cell(row=i+1, column=1).value = tubosdesconectados[i][0]
        ws3.cell(row=i+1, column=2).value = tubosdesconectados[i][1]
        ws3.cell(row=i+1, column=3).value = tubosdesconectados[i][2]
    
    # creo un estilo para las celdas

    font = Font(name='Calibri', size=11, bold=False, italic=False, vertAlign=None, underline='none', strike=False, color='FF000000')

    # recorro la hoja y le aplico el estilo a las celdas

    for row in ws3.iter_rows():
        for cell in row:
            cell.font = font
    
    # agrandar celdas de tablas
    for i in range(1,4):
        ws.column_dimensions[get_column_letter(i)].width = 30
        ws2.column_dimensions[get_column_letter(i)].width = 30
        ws3.column_dimensions[get_column_letter(i)].width = 30

    # guardo todo el excel

    wb.save(path)